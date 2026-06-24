# -*- coding: utf-8 -*-
"""Pasca-proses docx hasil Pandoc -> tata letak prosiding ITS yang rapi:
   1. Judul + identitas penulis 1 kolom; abstrak + isi 2 kolom.
  2. Semua gambar di-center; tabel dijaga agar tidak terpotong antar kolom.

Sesuai template Jurnal Teknik ITS: SEMUA gambar (termasuk diagram blok)
berada di dalam kolom (1 kolom), tidak ada gambar yang membentang dua kolom.
"""
import sys, copy
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL as VA

# =====================================================================
#  KONFIGURASI  — header (kop) gaya Jurnal Teknik ITS.
#  Teks di kiri, nomor halaman otomatis di kanan (sesuai TEMPLATE_PUBLIKASI_TEKNIK).
#  TODO: sesuaikan Vol/No/tahun bila perlu.
# =====================================================================
HEADER_TEXT = "JURNAL TEKNIK ITS Vol. X, No. Y, (TAHUN) ISSN: 2337-3539 (2301-9271 Print)"
DROPCAP_LINES = 3        # drop cap awal Pendahuluan setinggi N baris (utk huruf APA PUN)
# =====================================================================

doc = Document(sys.argv[1])
body = doc.element.body
final = body.find(qn('w:sectPr'))                # section terakhir (isi)

def style_by_name(name):
    """Cari style berdasarkan NAMA (bukan style_id) -> hindari DeprecationWarning
    python-docx 'style lookup by style_id is deprecated'. Kembalikan None bila
    tak ada, meniru perilaku KeyError pemanggil sebelumnya."""
    for st in doc.styles:
        if st.name == name:
            return st
    return None

def set_cols(sectPr, num, space=288):
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space))

def set_continuous(sectPr):
    t = sectPr.find(qn('w:type'))
    if t is None:
        t = OxmlElement('w:type'); sectPr.insert(0, t)
    t.set(qn('w:val'), 'continuous')

def add_sect(para, num):
    s = copy.deepcopy(final); set_cols(s, num); set_continuous(s)
    para._p.get_or_add_pPr().append(s)

def make_dropcap(body_para, lines=3, body_line_pt=10.35, x="947", y="135"):
    """Drop cap OTOMATIS: huruf pertama dibuat setinggi `lines` baris body, PERSIS,
    untuk HURUF APA PUN (L/M/W/I/dst). Tidak ada angka ukuran yang di-hardcode —
    semuanya dihitung dari tinggi 1 baris body (`body_line_pt`):

      - tinggi frame  w:h = lines x tinggiBaris  ->  teks SELALU membungkus TEPAT
        `lines` baris. Jumlah baris ditentukan tinggi frame, BUKAN hurufnya; lebar
        frame menyesuaikan lebar huruf secara otomatis (W lebar, I sempit).
      - ukuran huruf (sz), line-box (w:line), dan jarak atas (w:before) diskalakan
        dari tinggi frame memakai rasio hasil kalibrasi manual, supaya huruf MENGISI
        pas `lines` baris dan posisinya rata.

    Posisi frame (`x` = margin kiri, `y` = nudge atas) tetap. Bila ukuran/spasi font
    body berubah, cukup ubah `body_line_pt` (atau biarkan dihitung dari font) -> tetap
    pas N baris. Rasio diturunkan dari titik kalibrasi MANUAL di Word (finalisasi user):
    body 10,35pt -> frame 621, sz 84, line 615, before 132 (twip), x 947, y 135."""
    first_run = next((r for r in body_para.runs if r.text), None)
    if first_run is None:
        return None
    letter = first_run.text[0]
    first_run.text = first_run.text[1:]          # buang huruf pertama dari body

    line_tw  = body_line_pt * 20.0               # tinggi 1 baris (twip)
    frame_h  = round(lines * line_tw)            # N baris -> wrap N baris (kunci!)
    # ukuran huruf dikalibrasi dari setelan-tangan user di Word: kapital TNR yang
    # mengisi rapi N baris ~ frame_h/0,739/10 (rasio efektif kapital-ke-em hasil
    # finalisasi -> 621 twip menghasilkan sz 84 half-pt = 42pt).
    sz       = round(frame_h / 0.739 / 10.0)     # kapital setinggi ~N baris (isi penuh)
    szCs     = sz
    line_ex  = round(frame_h * 0.99)             # line-box huruf (exact) ~ tinggi frame
    before   = round(line_tw * 0.6377)           # jarak atas frame

    bpPr = body_para._p.find(qn('w:pPr'))
    ps = bpPr.find(qn('w:pStyle')) if bpPr is not None else None
    pstyle = '' if ps is None else f'<w:pStyle w:val="{ps.get(qn("w:val"))}"/>'
    rpr = (f'<w:rPr><w:sz w:val="{sz}"/><w:szCs w:val="{szCs}"/></w:rPr>')
    xml = (
        f'<w:p {nsdecls("w")}><w:pPr>{pstyle}'
        f'<w:framePr w:dropCap="drop" w:lines="{lines}" w:h="{frame_h}" '
        f'w:hRule="exact" w:wrap="around" w:vAnchor="text" w:hAnchor="page" '
        f'w:x="{x}" w:y="{y}"/>'
        f'<w:spacing w:before="{before}" w:line="{line_ex}" w:lineRule="exact"/>'
        f'{rpr}</w:pPr>'
        f'<w:r>{rpr}<w:t>{letter}</w:t></w:r></w:p>'
    )
    body_para._p.addprevious(parse_xml(xml))
    return letter

def force_black(style_name, no_underline=False):
    """Paksa warna teks suatu style -> hitam pekat (000000) & buang atribut theme.

    Heading 1/2/3 bawaan Pandoc memakai w:color themeColor="accent1" (biru tua
    0F4761) dan style Hyperlink biru 4F81BD; akibatnya judul subbab, e-mail, dan
    DOI/URL Daftar Pustaka tampil biru. Sama seperti gotcha font: bila atribut
    themeColor ADA, Word memprioritaskannya di atas w:val, jadi atribut theme
    DIBUANG dan w:val di-set 000000 agar benar-benar HITAM. Untuk Hyperlink,
    garis bawah dimatikan agar tampil seperti teks biasa.
    Dikerjakan di sini (pasca-proses, tiap build) supaya TIDAK menyentuh
    reference.docx yang sudah disetel-tangan -- lihat tools/style_reference.py."""
    st = style_by_name(style_name)
    if st is None:
        return
    rpr = st.element.get_or_add_rPr()
    color = rpr.find(qn('w:color'))
    if color is None:
        color = OxmlElement('w:color'); rpr.append(color)
    for a in ('w:themeColor', 'w:themeShade', 'w:themeTint'):
        if color.get(qn(a)) is not None:
            del color.attrib[qn(a)]
    color.set(qn('w:val'), '000000')
    if no_underline:
        u = rpr.find(qn('w:u'))
        if u is None:
            u = OxmlElement('w:u'); rpr.append(u)
        u.set(qn('w:val'), 'none')

def set_caption(style_name, sz="16"):
    """Style caption (Tabel/Gambar) -> 8pt (sz=16) & TIDAK bold.
    Pandoc memberi TableCaption/ImageCaption sz=18 (9pt) + <w:b/>; user minta
    8pt non-bold. jc dibiarkan apa adanya (both = rata kanan-kiri). Sama pola
    dgn force_black: dikerjakan di pasca-proses agar reference.docx tak disentuh."""
    st = style_by_name(style_name)
    if st is None:
        return
    rpr = st.element.get_or_add_rPr()
    for tag in ('w:sz', 'w:szCs'):
        e = rpr.find(qn(tag))
        if e is None:
            e = OxmlElement(tag); rpr.append(e)
        e.set(qn('w:val'), sz)
    for tag in ('w:b', 'w:bCs'):                       # buang bold
        e = rpr.find(qn(tag))
        if e is not None:
            rpr.remove(e)

# --- 0) warna teks heading + hyperlink -> HITAM (bukan biru theme) ---
for h in ("Heading 1", "Heading 2", "Heading 3",
          "Heading 1 Char", "Heading 2 Char", "Heading 3 Char"):
    force_black(h)
force_black("Hyperlink", no_underline=True)
print("Warna heading + hyperlink -> hitam (000000).")

# --- 0a) samakan format dgn TEMPLATE_PUBLIKASI_TEKNIK (pasca-proses) ---
# Heading 1: spasi sebelum/sesudah 18/4pt. (Teks judul bab diketik KAPITAL
# langsung di bagian/*.tex, mengikuti cara template resmi.)
st_h1 = style_by_name("Heading 1")
if st_h1 is not None:
    # space-before = 0 supaya heading di PUNCAK kolom SEJAJAR dgn kolom sebelah.
    # 'Air' di atas heading diberikan via space-after paragraf sebelumnya (langkah
    # 7) -> di tengah kolom tetap lega, di puncak kolom otomatis ditekan Word.
    st_h1.paragraph_format.space_before = Pt(0)
    st_h1.paragraph_format.space_after = Pt(4)
    print("Heading 1 -> space-before 0 (sejajar), space-after 4.")
# Blok Penulis: 11pt (template: nama penulis 11pt).
st_au = style_by_name("Author")
if st_au is not None:
    st_au.font.size = Pt(11)
    print("Penulis -> 11pt (sesuai template).")

# --- 0b) caption Tabel/Gambar -> 8pt & non-bold ---
for c in ("Table Caption", "Image Caption", "Caption"):
    set_caption(c)
print("Caption Tabel/Gambar -> 8pt, non-bold.")

# --- isi = 2 kolom ---
set_cols(final, 2, 288); set_continuous(final)

paras = doc.paragraphs

def idx_startswith(prefix):
    for i, p in enumerate(paras):
        if p.text.strip().startswith(prefix):
            return i
    return None

# --- 1) front-matter 1 kolom + jarak FULL-WIDTH di bawah e-mail ---
# Judul/penulis/e-mail 1 kolom; Abstrak dst 2 kolom. Agar KEDUA kolom mulai
# sejajar dengan jarak yang sama di bawah e-mail (seperti template), paragraf
# jarak HARUS berada di section 1 (full width), BUKAN di kolom kiri saja.
# Caranya: section-break 1 kolom dipasang pada paragraf-kosong tepat SEBELUM
# Abstrak -> paragraf itu jadi penutup section 1 (full width) sekaligus jaraknya.
# (Versi lama menaruh paragraf kosong di section 2 sehingga hanya kolom kiri
#  yang turun, kolom kanan tetap mepet ke e-mail.)
FRONT_GAP = Pt(12)            # jarak di bawah e-mail (atur di sini bila perlu)
i_email = None
for i, p in enumerate(paras):
    if 'e-mail:' in p.text or 'student.its.ac.id' in p.text:
        i_email = i
        break
i_abs = idx_startswith("Abstrak")
if i_email is not None:
    # Section 1 (1 kolom) berakhir di paragraf e-mail. JARAK di bawah e-mail
    # dibuat full-width via space-after paragraf e-mail itu sendiri -> KEDUA
    # kolom section 2 (Abstrak dst) mulai SEJAJAR di bawah jarak yang sama.
    # (Paragraf-kosong tidak dipakai: bila ia memuat section-break, tingginya
    #  runtuh dan jaraknya hilang.)
    add_sect(paras[i_email], 1)
    paras[i_email].paragraph_format.space_after = FRONT_GAP
    print("Front-matter 1 kolom: break + jarak full-width (space-after e-mail).")
elif i_abs is not None:
    gap = paras[i_abs].insert_paragraph_before("")    # fallback: break sebelum Abstrak
    add_sect(gap, 1)
    print("Front-matter 1 kolom: break sebelum Abstrak (fallback).")
else:
    i_kw = idx_startswith("Kata Kunci")               # fallback: break di Kata Kunci
    if i_kw is not None:
        add_sect(paras[i_kw], 1)
        print("Front-matter 1 kolom: break setelah Kata Kunci (fallback).")

# --- 1c) drop cap huruf pertama Pendahuluan (estetika, seperti template) ---
def first_body_after_heading(keyword):
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        if p.style.name.startswith("Heading") and keyword.lower() in p.text.lower():
            for j in range(i + 1, len(ps)):
                if ps[j].text.strip():
                    return ps[j]
    return None
bp = first_body_after_heading("Pendahuluan")
if bp is not None:
    # tinggi 1 baris body = ukuran font body x rasio spasi-tunggal TNR (~1,035).
    # Dihitung dari font -> ikut menyesuaikan bila ukuran font body diubah.
    _bt = style_by_name("Body Text") or style_by_name("Normal")
    _body_pt = _bt.font.size.pt if (_bt is not None and _bt.font.size is not None) else 10.0
    _line_pt = _body_pt * 1.035
    huruf = make_dropcap(bp, lines=DROPCAP_LINES, body_line_pt=_line_pt)
    print(f"Drop cap '{huruf}' otomatis: pas {DROPCAP_LINES} baris "
          f"(body {_body_pt:.1f}pt -> baris {_line_pt:.2f}pt).")

# --- 1d) seluruh ABSTRAK + baris KATA KUNCI dicetak bold (gaya template ITS) ---
# Label "Abstrak"/"Kata Kunci" sudah bold dari .tex; di sini SEMUA run di paragraf
# abstrak (style "Abstract") dan paragraf kata kunci ikut di-bold. Istilah italic
# tetap miring -> jadi bold-italic. Dikerjakan di pasca-proses agar tak menyentuh
# style "FirstParagraph" yg dipakai paragraf lain.
n_bold = 0
for p in doc.paragraphs:
    pPr = p._p.find(qn('w:pPr'))
    ps = pPr.find(qn('w:pStyle')) if pPr is not None else None
    sid = ps.get(qn('w:val')) if ps is not None else ''
    is_kw = p.text.strip().startswith('Kata Kunci')
    if sid == 'Abstract' or is_kw:
        for r in p.runs:
            r.font.bold = True
            if is_kw:
                r.font.size = Pt(9)            # Kata Kunci 9pt (samai Abstrak/template)
        if p.runs:
            n_bold += 1
print("Abstrak + Kata Kunci di-bold (paragraf):", n_bold)

# --- 1e) alinea (first-line indent) badan teks, sesuai TEMPLATE_PUBLIKASI_TEKNIK ---
# Template resmi: style "Text" (badan) ber-first-line-indent 10,1pt (=202 twip ~0,36cm);
# Judul/Penulis/Heading/Persamaan = 0, Daftar Pustaka = hanging indent.
# Pandoc memetakan badan teks ke style BodyText + FirstParagraph (paragraf pertama
# setelah heading). Keduanya kita beri indent, KECUALI:
#   - baris "Kata Kunci" (template: tanpa indent),
#   - Abstrak (style "Abstract" — paragraf pertama abstrak template tanpa indent),
#   - paragraf drop-cap Pendahuluan: huruf pertama (ber-framePr) + badan yg membungkusnya,
#   - Daftar Pustaka (style Bibliography, hanging indent dibiarkan apa adanya).
INDENT = Pt(10.1)
n_indent = 0
for p in doc.paragraphs:
    pPr = p._p.find(qn('w:pPr'))
    ps = pPr.find(qn('w:pStyle')) if pPr is not None else None
    sid = ps.get(qn('w:val')) if ps is not None else ''
    if sid not in ('BodyText', 'FirstParagraph'):
        continue
    if p.text.strip().startswith('Kata Kunci'):
        continue
    if pPr is not None and pPr.find(qn('w:framePr')) is not None:
        continue                                   # paragraf drop-cap
    if bp is not None and p._p is bp._p:
        continue                                   # badan yg membungkus drop-cap
    p.paragraph_format.first_line_indent = INDENT
    n_indent += 1
print("Alinea (first-line indent 10,1pt) dipasang:", n_indent)

# --- 2) center semua paragraf bergambar (termasuk gambar di dalam sel tabel) ---
# Catatan: gambar dengan DUA includegraphics dalam satu figure (mis. Gambar a+b)
# dibungkus Pandoc jadi tabel 1-sel; paragrafnya tak terjangkau
# doc.paragraphs, jadi kita iterasi SEMUA <w:p>.
n_img = 0
for p in doc.element.body.iter(qn('w:p')):
    if p.find('.//' + qn('w:drawing')) is not None:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc'); pPr.append(jc)
        jc.set(qn('w:val'), 'center')
        n_img += 1
print("Gambar di-center:", n_img)

# --- 2b) tabel-pembungkus gambar (1 sel, berisi drawing): center + rapat ---
for tbl in doc.tables:
    if tbl._tbl.find('.//' + qn('w:drawing')) is not None and len(tbl.rows) == 1:
        jc = tbl._tbl.find(qn('w:tblPr') + '/' + qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc'); tbl._tbl.find(qn('w:tblPr')).append(jc)
        jc.set(qn('w:val'), 'center')

# --- 4) tabel: cegah baris terpotong + tabel utuh sebagai satu blok ---
for tbl in doc.tables:
    rows = tbl.rows
    n = len(rows)
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))   # baris tak terbelah
        if i < n - 1:                                  # keep-with-next antar-baris
            for cell in row.cells:                     # -> tabel tak terpotong
                for cp in cell.paragraphs:             #    antar kolom/halaman
                    cp.paragraph_format.keep_with_next = True
for p in doc.paragraphs:
    if p.text.strip().startswith("Tabel "):
        p.paragraph_format.keep_with_next = True       # caption nempel ke tabel

# --- 4b) isi sel tabel data: rata tengah (horizontal + vertikal) biar rapi ---
# Semua tabel data dibuat seragam center; tabel pembungkus gambar (berisi
# drawing) dilewati supaya penataan gambarnya tidak terganggu.
n_celltbl = 0
for tbl in doc.tables:
    if tbl._tbl.find('.//' + qn('w:drawing')) is not None:
        continue
    # lebar tabel = 100% kolom (rata batas kanan-kiri kolom), kolom diskala
    # proporsional; sel tak punya tcW sendiri jadi cukup set tblW=pct 5000.
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct'); tblW.set(qn('w:w'), '5000')
    tbl.autofit = True
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = VA.CENTER
            for cp in cell.paragraphs:
                cp.alignment = AL.CENTER
    n_celltbl += 1
print("Tabel data: rata-tengah + lebar 100% kolom:", n_celltbl)

# --- 5) paksa A4 + margin ITS + header semua section ---
# Header prosiding (3 baris, TNR 10pt, rata kiri pojok atas). Jarak header dari
# atas & footer dari bawah = 0,76 cm. Diset di SEMUA section dgn isi identik agar
# muncul sama di setiap halaman (termasuk halaman dgn section break kontinu).
def _tnr10(run):                                        # paksa TNR 10pt (hindari theme font)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), "Times New Roman")

def _add_page_field(run):                               # sisipkan field { PAGE } -> nomor halaman
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(instr); run._r.append(end)

def fill_header(hdr):
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    for r in list(p.runs):                              # kosongkan paragraf header
        r._element.getparent().remove(r._element)
    p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0               # single (alami, tdk rapat)
    # tab kanan di batas kolom (lebar = 21 - 1,65 - 1,65 = 17,7 cm) untuk nomor halaman
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.7), WD_TAB_ALIGNMENT.RIGHT)
    _tnr10(p.add_run(HEADER_TEXT))                      # teks jurnal (kiri)
    rt = p.add_run(); rt.add_tab(); _tnr10(rt)          # tab -> dorong ke kanan
    rp = p.add_run(); _tnr10(rp); _add_page_field(rp)   # nomor halaman (field PAGE)

for sec in doc.sections:
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(1.65)
    # Margin = TEMPLATE_PUBLIKASI_TEKNIK: atas/bawah 1,78 cm, kiri/kanan 1,65 cm.
    sec.top_margin = Cm(1.78)
    sec.bottom_margin = Cm(1.78)
    sec.header_distance = Cm(0.76)
    sec.footer_distance = Cm(0.76)
    sec.different_first_page_header_footer = False
    fill_header(sec.header)
print("Header Jurnal Teknik ITS (TNR 10pt) + nomor halaman dipasang di semua section.")

# --- 6) slot LAMPIRAN di paling akhir (setelah Daftar Pustaka) ---
# Heading kosong sebagai tempat lampiran; isinya menyusul (boleh dibiarkan kosong).
# Daftar Pustaka dibangkitkan Pandoc sebagai konten terakhir body, jadi paragraf
# baru di sini otomatis jatuh SETELAH-nya (sebelum sectPr akhir).
lampiran = doc.add_paragraph("LAMPIRAN", style="Heading 1")
lampiran.alignment = AL.CENTER                         # seragam dgn DAFTAR PUSTAKA
doc.add_paragraph("")                                  # ruang isi lampiran
print("Slot LAMPIRAN ditambahkan setelah Daftar Pustaka.")

# --- 7) 'air' sebelum tiap Heading 1 = space-AFTER paragraf SEBELUMnya ---
# Agar heading di tengah kolom tetap lega TAPI heading di puncak kolom/halaman
# tetap SEJAJAR: ruang di atas heading tidak dipasang sebagai space-before heading
# (yang ikut tampil di puncak kolom & merusak kesejajaran), melainkan sebagai
# space-after paragraf sebelumnya. Word otomatis menekan space di DASAR kolom,
# sehingga saat heading pindah ke puncak kolom ruangnya hilang -> sejajar; saat di
# tengah kolom ruangnya tampil -> lega. (suppressSpBfAfterPgBrk TIDAK dipakai
# karena hanya berlaku utk page-break, bukan column-break.)
HEAD_AIR = Pt(16)
_ap = doc.paragraphs
for i, p in enumerate(_ap):
    if i > 0 and p.style.name == "Heading 1":
        _ap[i - 1].paragraph_format.space_after = HEAD_AIR
print("Air sebelum Heading 1 -> space-after paragraf sebelumnya (sejajar + lega).")

doc.save(sys.argv[1])
print("Layout selesai. Jumlah section:", len(doc.sections))
